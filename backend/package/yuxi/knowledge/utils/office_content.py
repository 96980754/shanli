"""Word/Excel 内容提取：将 docx/xlsx 解析为可编辑的结构化内容。

- .docx → blocks: [{kind: 'heading'|'para'|'table', text|rows}]
- .xlsx → sheets: [{name, rows: [[cell, ...]]}]

只提取文字与单元格内容，不保留复杂格式，供网页端编辑后写回。
"""

from __future__ import annotations

import os


def _docx_blocks(file_path: str) -> list[dict]:
    from docx import Document

    document = Document(file_path)
    blocks: list[dict] = []

    # 按文档顺序遍历段落与表格，保持阅读顺序
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            para_text = child.text or ""
            # python-docx 直接访问段落文本更可靠
            text = _para_text(document, child)
            if not text:
                continue
            # 依据样式粗判标题（Heading 1-6）
            style_name = _para_style(document, child)
            kind = "heading" if style_name and style_name.lower().startswith("heading") else "para"
            blocks.append({"kind": kind, "text": text})
        elif tag == "tbl":
            rows = _table_rows(document, child)
            if rows:
                blocks.append({"kind": "table", "rows": rows})

    # 兜底：若上面遍历无输出，退回到简单遍历
    if not blocks:
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ""
                kind = "heading" if style and style.lower().startswith("heading") else "para"
                blocks.append({"kind": kind, "text": text})
        for table in document.tables:
            rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
            if rows:
                blocks.append({"kind": "table", "rows": rows})

    return blocks


def _para_text(document, p) -> str:
    """从 lxml 段落元素提取合并文本。"""
    texts = [t.text or "" for t in p.iter() if t.tag.rsplit("}", 1)[-1] == "t"]
    return "".join(texts).strip()


def _para_style(document, p) -> str | None:
    try:
        from docx.text.paragraph import Paragraph

        return Paragraph(p, document).style.name if p is not None else None
    except Exception:
        return None


def _table_rows(document, tbl) -> list[list[str]]:
    from docx.table import Table

    try:
        table = Table(tbl, document)
    except Exception:
        return []
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(cells):
            rows.append(cells)
    return rows


def _xlsx_sheets(file_path: str) -> list[dict]:
    from openpyxl import load_workbook

    workbook = load_workbook(file_path, data_only=True, read_only=True)
    sheets: list[dict] = []
    for sheet in workbook.worksheets:
        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if cell is None else str(cell).strip() for cell in row]
            if any(cells):
                rows.append(cells)
        sheets.append({"name": sheet.title, "rows": rows})
    return sheets


async def extract_office_content(kb_id: str, file_path: str, filename: str) -> dict:
    """从 MinIO 读取并解析 Word/Excel 为可编辑结构。

    Args:
        kb_id: 知识库 ID
        file_path: MinIO 原始文件 URL
        filename: 原始文件名（决定 docx/xlsx）

    Returns:
        dict: {type: 'docx'|'xlsx', blocks|sheets}
    """
    import tempfile

    from yuxi.knowledge.runtime import knowledge_base

    kb_instance = await knowledge_base._get_kb_for_database(kb_id)
    raw_bytes = await kb_instance._read_minio_bytes(file_path)
    suffix = os.path.splitext(filename or "")[1].lower() or ".bin"
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(raw_bytes)
            temp_path = tmp.name
        if suffix == ".xlsx":
            return {"type": "xlsx", "sheets": _xlsx_sheets(temp_path)}
        # .docx 及未知按 docx 处理
        return {"type": "docx", "blocks": _docx_blocks(temp_path)}
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except OSError:
                pass
