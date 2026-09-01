"""Word/Excel 写回：把网页编辑后的文字/单元格写回 docx/xlsx。

- write_docx(blocks) → bytes：python-docx 生成 .docx
- write_xlsx(sheets) → bytes：openpyxl 生成 .xlsx
- markdown_to_blocks(markdown) / markdown_to_sheets(markdown)：清洗后的 markdown 转结构化内容，供写回原格式
- replace_document_content(...)：写回后入库并删旧版
"""

from __future__ import annotations

import io
import os
from typing import Any

from markdown_it import MarkdownIt

# 与清洗链路共用同一 markdown 解析器（document_section_splitter 同款），
# 保证清洗后的 markdown 写回原格式时标题/表格解析一致。
_md = MarkdownIt("commonmark", {"html": True}).enable("table")

# 白皮书品牌信息：封面 logo / 公司名 / 版权与免责声明，对齐甲方（POCSTARS）白皮书案例。
_WHITEPAPER_COMPANY = "善理通益信息科技（深圳）有限公司"
_WHITEPAPER_LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "pocstars_logo.png")
_WHITEPAPER_COPYRIGHT = (
    "POCSTARS为善理通益信息科技（深圳）有限公司（以下简称“我司”）在中华人民共和国和世界其他国家和/或地区的商标或注册商标。"
    "我司仅对自有商标及产品名称享有所有权，本文中可能提及的其他商标和产品名称为各自所有者所有。",
    "本文介绍的产品中可能包括存储于内存或其他媒介中的计算机程序。我司对此等程序享有的专有权利受中华人民共和国或其他国家及相关国际法的保护。"
    "购买本产品并不意味着我司以明示或暗示方式向购买者授予有关此等电脑程序的权益。未经我司事先书面授权，任何企业、组织或个人不得对计算机程序进行任何形式的复制、更改、散发、反编译和反向工程。",
)
_WHITEPAPER_DISCLAIMER = (
    "本文在编制过程中力求内容的准确性与完整性，但对于可能出现的错误或疏漏，我司不承担任何责任。由于技术的不断发展，我司保留不予通知而更改产品设计与规格的权利。"
    "未经我司事先书面授权，不得以任何形式对本文进行复制、修改、翻译和散发。",
    "如需更多信息或对本文有任何建议，欢迎访问我们的网站：https://www.pocstars.cn/。",
)
_WHITEPAPER_AI_NOTE = "请注意，本内容由AI生成。"


def _set_run_font(run: Any, font_name: str) -> None:
    from docx.oxml.ns import qn

    run.font.name = font_name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), font_name)


def _configure_docx_styles(document: Any, font_name: str) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style_specs = {
        "Normal": (11, False),
        "Title": (20, True),
        "Heading 1": (18, True),
        "Heading 2": (15, True),
        "Heading 3": (13, True),
        "Heading 4": (11, True),
        "Heading 5": (11, True),
        "Heading 6": (11, True),
        "List Bullet": (11, False),
        "List Number": (11, False),
    }
    for style_name, (size, bold) in style_specs.items():
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "eastAsia"):
            fonts.set(qn(f"w:{attribute}"), font_name)


def _write_inline(paragraph: Any, value: Any, font_name: str | None) -> None:
    if isinstance(value, dict):
        runs = value.get("runs")
        text = str(value.get("text") or "")
    else:
        runs = None
        text = str(value or "")

    if runs:
        for item in runs:
            run = paragraph.add_run(str(item.get("text") or ""))
            run.bold = bool(item.get("bold"))
            run.italic = bool(item.get("italic"))
            if font_name:
                _set_run_font(run, font_name)
        return

    run = paragraph.add_run(text)
    if font_name:
        _set_run_font(run, font_name)


def _write_blocks(document: Any, blocks: list[dict] | None, font_name: str | None) -> None:
    """把 markdown_to_blocks 得到的 blocks 写入已创建的 document。"""
    for block in blocks or []:
        kind = block.get("kind")
        if kind == "heading":
            level = min(max(int(block.get("level") or 1), 1), 6)
            paragraph = document.add_heading(level=level)
            _write_inline(paragraph, block, font_name)
        elif kind == "table":
            rows = block.get("rows") or []
            if rows:
                ncols = max(len(r) for r in rows)
                table = document.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j in range(ncols):
                        cell = table.cell(i, j)
                        _write_inline(cell.paragraphs[0], row[j] if j < len(row) else "", font_name)
        elif kind == "list_item":
            paragraph = document.add_paragraph(style="List Number" if block.get("ordered") else "List Bullet")
            _write_inline(paragraph, block, font_name)
        else:
            paragraph = document.add_paragraph()
            _write_inline(paragraph, block, font_name)


def write_docx(blocks: list[dict], *, font_name: str | None = None) -> bytes:
    """将编辑后的段落/标题/表格写回 .docx 字节。"""
    from docx import Document

    document = Document()
    if font_name:
        _configure_docx_styles(document, font_name)
    _write_blocks(document, blocks, font_name)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


_WHITEPAPER_ACCENT_HEX = "333399"


def write_whitepaper_docx(
    blocks: list[dict],
    *,
    cover: dict[str, str],
    header_text: str,
    font_name: str | None = None,
) -> bytes:
    """生成白皮书级排版的 .docx：封面 + 目录域 + 页眉页脚 + 正文块。

    专用于行业解决方案等对外交付文档；文档编辑写回仍走 write_docx，互不影响。
    """
    from docx import Document
    from docx.enum.section import WD_SECTION

    document = Document()
    if font_name:
        _configure_docx_styles(document, font_name)
    _configure_whitepaper_styles(document)
    _configure_whitepaper_page(document.sections[0])

    _write_cover(document, cover, font_name)
    document.add_page_break()
    _write_legal_page(document, font_name)
    document.add_section(WD_SECTION.NEW_PAGE)

    body_section = document.sections[-1]
    _configure_whitepaper_header_footer(body_section, header_text, font_name)

    _write_toc(document, font_name)
    document.add_page_break()
    _write_blocks(document, blocks, font_name)
    _style_solution_tables(document)
    _write_ai_disclaimer(document, font_name)
    _enable_automatic_field_update(document)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _configure_whitepaper_styles(document: Any) -> None:
    """白皮书强调色：章节标题用深蓝紫 #333399，并加大段前段后间距。"""
    from docx.shared import Pt, RGBColor

    accent = RGBColor.from_string(_WHITEPAPER_ACCENT_HEX)
    for name, (size, before, after) in {
        "Heading 1": (Pt(18), Pt(18), Pt(8)),
        "Heading 2": (Pt(15), Pt(12), Pt(6)),
        "Heading 3": (Pt(13), Pt(10), Pt(4)),
        "Heading 4": (Pt(11), Pt(8), Pt(4)),
    }.items():
        style = document.styles[name]
        style.font.size = size
        style.font.color.rgb = accent
        style.paragraph_format.space_before = before
        style.paragraph_format.space_after = after


def _configure_whitepaper_page(section: Any) -> None:
    """对齐甲方白皮书案例：A4 纸型 + 左右 1.81cm、上下 2.00cm 边距。

    新增节会继承本节页面设置，故只配置首节即可覆盖封面与正文两节。
    """
    from docx.shared import Cm

    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.81)
    section.right_margin = Cm(1.81)
    section.top_margin = Cm(2.00)
    section.bottom_margin = Cm(2.00)


def _write_cover(document: Any, cover: dict[str, str], font_name: str | None) -> None:
    """封面：品牌 logo + 标签 + 大标题 + 行业/产品 + 编制日期 + 公司名，全部居中。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    accent = RGBColor.from_string(_WHITEPAPER_ACCENT_HEX)
    dark = RGBColor(0x1F, 0x1F, 0x1F)
    gray = RGBColor(0x40, 0x40, 0x40)
    light_gray = RGBColor(0x66, 0x66, 0x66)

    def _center(text: str, size: int, color: RGBColor, *, bold: bool = False) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color
        if font_name:
            _set_run_font(run, font_name)

    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.paragraph_format.space_after = Pt(24)
    logo_paragraph.add_run().add_picture(_WHITEPAPER_LOGO_PATH, width=Cm(5.5))

    for _ in range(4):
        document.add_paragraph()
    label = str(cover.get("label") or "").strip()
    if label:
        _center(label, 15, accent, bold=True)
        document.add_paragraph()
    title = str(cover.get("title") or "").strip()
    if title:
        _center(title, 30, dark, bold=True)
        document.add_paragraph()
    industry = str(cover.get("industry") or "").strip()
    if industry:
        _center(f"行业 / 场景：{industry}", 14, gray)
        document.add_paragraph()
    products = str(cover.get("products") or "").strip()
    if products:
        _center(f"选用产品：{products}", 14, gray)
    for _ in range(6):
        document.add_paragraph()
    date_text = str(cover.get("date") or "").strip()
    if date_text:
        _center(date_text, 11, light_gray)
        document.add_paragraph()
    _center(_WHITEPAPER_COMPANY, 11, light_gray)


def _write_legal_page(document: Any, font_name: str | None) -> None:
    """版权信息 + 免责声明页（对齐甲方案例；普通段落样式，避免进入目录）。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    accent = RGBColor.from_string(_WHITEPAPER_ACCENT_HEX)
    gray = RGBColor(0x40, 0x40, 0x40)

    def _heading(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(text)
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = accent
        if font_name:
            _set_run_font(run, font_name)

    def _body(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.4
        run = paragraph.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = gray
        if font_name:
            _set_run_font(run, font_name)

    _heading("版权信息")
    for paragraph_text in _WHITEPAPER_COPYRIGHT:
        _body(paragraph_text)
    _heading("免责声明")
    for paragraph_text in _WHITEPAPER_DISCLAIMER:
        _body(paragraph_text)


def _write_ai_disclaimer(document: Any, font_name: str | None) -> None:
    """文档末尾显著提示：本内容由 AI 生成（红色加粗 + 上下红色边框）。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(30)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "12")
        element.set(qn("w:space"), "6")
        element.set(qn("w:color"), "C00000")
        p_bdr.append(element)
    p_pr.append(p_bdr)
    run = paragraph.add_run(_WHITEPAPER_AI_NOTE)
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    if font_name:
        _set_run_font(run, font_name)


def _write_toc(document: Any, font_name: str | None) -> None:
    """目录标题（非 Heading 样式，避免自引用进目录）+ 自动更新的 TOC 域。"""
    from docx.shared import Pt, RGBColor

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("目  录")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(_WHITEPAPER_ACCENT_HEX)
    if font_name:
        _set_run_font(run, font_name)
    _append_field_run(
        document.add_paragraph(),
        'TOC \\o "1-2" \\h \\z \\u',
        placeholder="目录将在打开文档时自动更新",
        font_name=font_name,
    )


def _configure_whitepaper_header_footer(section: Any, header_text: str, font_name: str | None) -> None:
    """正文节页眉：方案标题 + 底部强调色分隔线；页脚：居中「第 X 页」页码域。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_paragraph.add_run(header_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if font_name:
        _set_run_font(run, font_name)
    _add_paragraph_bottom_border(header_paragraph)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_field_run(
        footer_paragraph,
        " PAGE ",
        placeholder="1",
        prefix="第 ",
        suffix=" 页",
        font_name=font_name,
    )


def _append_field_run(
    paragraph: Any,
    instruction: str,
    *,
    placeholder: str = "",
    prefix: str = "",
    suffix: str = "",
    font_name: str | None = None,
) -> None:
    """在段落末尾追加一个 Word 域（begin / instrText / separate / 占位 / end）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _fld_char(char_type: str) -> OxmlElement:
        element = OxmlElement("w:fldChar")
        element.set(qn("w:fldCharType"), char_type)
        return element

    if prefix:
        paragraph.add_run(prefix)
    run = paragraph.add_run()
    run._r.append(_fld_char("begin"))
    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run._r.append(instr)
    run = paragraph.add_run()
    run._r.append(_fld_char("separate"))
    if placeholder:
        run = paragraph.add_run(placeholder)
        if font_name:
            _set_run_font(run, font_name)
    run = paragraph.add_run()
    run._r.append(_fld_char("end"))
    if suffix:
        run = paragraph.add_run(suffix)
        if font_name:
            _set_run_font(run, font_name)


def _add_paragraph_bottom_border(paragraph: Any, *, color: str = _WHITEPAPER_ACCENT_HEX, size: str = "6") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _style_solution_tables(document: Any) -> None:
    """表格白皮书样式：细边框、表头 #333399 填充 + 白字加粗 + 跨页重复表头。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    for table in document.tables:
        tbl_pr = table._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:color"), "BFBFBF")
            borders.append(element)
        tbl_pr.append(borders)

        cell_margin = OxmlElement("w:tblCellMar")
        for side, width in (("top", 40), ("left", 80), ("bottom", 40), ("right", 80)):
            element = OxmlElement(f"w:{side}")
            element.set(qn("w:w"), str(width))
            element.set(qn("w:type"), "dxa")
            cell_margin.append(element)
        tbl_pr.append(cell_margin)

        if not table.rows:
            continue
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), _WHITEPAPER_ACCENT_HEX)
            tc_pr.append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _enable_automatic_field_update(document: Any) -> None:
    """让 Word 打开文档时自动更新 TOC 与页码等域。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    document.settings.element.append(update_fields)


def write_xlsx(sheets: list[dict]) -> bytes:
    """将编辑后的工作表写回 .xlsx 字节。"""
    from openpyxl import Workbook

    workbook = Workbook()
    for idx, sheet in enumerate(sheets or []):
        name = (sheet.get("name") or f"Sheet{idx + 1}")[:31]
        ws = workbook.active if idx == 0 else workbook.create_sheet()
        ws.title = name
        for row in sheet.get("rows") or []:
            ws.append(list(row))
    buf = io.BytesIO()
    workbook.save(buf)
    return buf.getvalue()


def _inline_value(token: Any) -> str | dict[str, Any]:
    children = token.children or []
    if not children:
        return str(token.content or "").strip()

    runs: list[dict[str, Any]] = []
    bold_depth = 0
    italic_depth = 0
    for child in children:
        if child.type == "strong_open":
            bold_depth += 1
            continue
        if child.type == "strong_close":
            bold_depth = max(0, bold_depth - 1)
            continue
        if child.type == "em_open":
            italic_depth += 1
            continue
        if child.type == "em_close":
            italic_depth = max(0, italic_depth - 1)
            continue
        if child.type in {"softbreak", "hardbreak"}:
            text = "\n"
        elif child.type in {"text", "code_inline", "html_inline"}:
            text = str(child.content or "")
        else:
            continue
        if not text:
            continue
        item = {"text": text, "bold": bold_depth > 0, "italic": italic_depth > 0}
        if runs and all(runs[-1][key] == item[key] for key in ("bold", "italic")):
            runs[-1]["text"] += text
        else:
            runs.append(item)

    text = "".join(item["text"] for item in runs).strip()
    if not text:
        return ""
    if not any(item["bold"] or item["italic"] for item in runs):
        return text
    return {"text": text, "runs": runs}


def _collect_table_rows(tokens: list[Any], start: int) -> list[list[Any]]:
    """从 table_open 开始收集表格行，每行一个单元格字符串列表。"""
    rows: list[list[Any]] = []
    current: list[Any] = []
    j = start + 1
    n = len(tokens)
    while j < n and tokens[j].type != "table_close":
        if tokens[j].type == "tr_open":
            current = []
        elif tokens[j].type == "tr_close":
            if current:
                rows.append(current)
                current = []
        elif tokens[j].type == "inline":
            value = _inline_value(tokens[j])
            if value:
                current.append(value)
        j += 1
    if current:
        rows.append(current)
    return rows


def _skip_to_token(tokens: list[Any], start: int, token_type: str) -> int:
    """返回从 start 起第一个 type == token_type 的下标，找不到则返回 len。"""
    n = len(tokens)
    j = start
    while j < n and tokens[j].type != token_type:
        j += 1
    return j


def markdown_to_blocks(markdown: str) -> list[dict]:
    """将清洗后的 markdown 解析为 docx blocks: [{kind: heading|para|table, text|rows, level?}]"""
    if not markdown or not markdown.strip():
        return []
    tokens = _md.parse(markdown)
    blocks: list[dict] = []
    list_stack: list[bool] = []
    i = 0
    n = len(tokens)
    while i < n:
        ttype = tokens[i].type
        if ttype == "bullet_list_open":
            list_stack.append(False)
            i += 1
        elif ttype == "ordered_list_open":
            list_stack.append(True)
            i += 1
        elif ttype in {"bullet_list_close", "ordered_list_close"}:
            if list_stack:
                list_stack.pop()
            i += 1
        elif ttype == "heading_open":
            level = int(tokens[i].tag[1:]) if tokens[i].tag and len(tokens[i].tag) > 1 else 1
            inline = tokens[i + 1] if i + 1 < n and tokens[i + 1].type == "inline" else None
            value = _inline_value(inline) if inline else ""
            text = value.get("text", "") if isinstance(value, dict) else value
            if text:
                block = {"kind": "heading", "text": text, "level": level}
                if isinstance(value, dict):
                    block["runs"] = value["runs"]
                blocks.append(block)
            i += 3
        elif ttype == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < n and tokens[i + 1].type == "inline" else None
            value = _inline_value(inline) if inline else ""
            text = value.get("text", "") if isinstance(value, dict) else value
            if text:
                block = {
                    "kind": "list_item" if list_stack else "para",
                    "text": text,
                }
                if list_stack:
                    block["ordered"] = list_stack[-1]
                if isinstance(value, dict):
                    block["runs"] = value["runs"]
                blocks.append(block)
            i += 3
        elif ttype == "table_open":
            rows = _collect_table_rows(tokens, i)
            if rows:
                blocks.append({"kind": "table", "rows": rows})
            i = _skip_to_token(tokens, i, "table_close") + 1
        else:
            i += 1
    return blocks


def markdown_to_sheets(markdown: str) -> list[dict]:
    """将清洗后的 markdown 解析为 xlsx sheets: [{name, rows: [[...]]}]。

    表格行保留多列结构；标题/段落作为单列文本行，全部归入默认 Sheet。
    """
    if not markdown or not markdown.strip():
        return []
    tokens = _md.parse(markdown)
    rows: list[list[str]] = []
    i = 0
    n = len(tokens)
    while i < n:
        ttype = tokens[i].type
        if ttype == "heading_open":
            inline = tokens[i + 1] if i + 1 < n and tokens[i + 1].type == "inline" else None
            text = (inline.content or "").strip() if inline else ""
            if text:
                rows.append([text])
            i += 3
        elif ttype == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < n and tokens[i + 1].type == "inline" else None
            text = (inline.content or "").strip() if inline else ""
            if text:
                rows.append([text])
            i += 3
        elif ttype == "table_open":
            rows.extend(
                [
                    [cell.get("text", "") if isinstance(cell, dict) else cell for cell in row]
                    for row in _collect_table_rows(tokens, i)
                ]
            )
            i = _skip_to_token(tokens, i, "table_close") + 1
        else:
            i += 1
    return [{"name": "Sheet1", "rows": rows}] if rows else []


def serialize_edited_content(content_type: str, payload: dict) -> bytes:
    """根据 type 把编辑后的 blocks/sheets 序列化为 docx/xlsx 字节。"""
    if content_type == "xlsx":
        return write_xlsx(payload.get("sheets") or [])
    return write_docx(payload.get("blocks") or [])
