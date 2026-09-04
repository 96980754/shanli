from __future__ import annotations

import io
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from docx.oxml.ns import qn
from pydantic import ValidationError

import yuxi.agents.toolkits.industry_solution as industry_tools
from server.routers.agent_router import AgentRunCreate
from yuxi.services.industry_solution_service import (
    IndustrySolutionRequest,
    render_solution_docx,
    sanitize_docx_filename,
)
from yuxi.services.input_message_service import build_chat_input_message


def test_industry_solution_tools_register_runtime_as_injected_argument():
    assert "runtime" in industry_tools.research_industry_products._injected_args_keys
    assert "runtime" in industry_tools.export_industry_solution_docx._injected_args_keys


@pytest.mark.asyncio
async def test_research_industry_products_retrieves_each_product_and_keeps_sources(monkeypatch):
    queries: list[str] = []

    async def fake_retrieve(kb_ids, query_text, *, runtime, file_name=None):
        queries.append(query_text)
        product = "产品A" if "产品A" in query_text else "产品B"
        other_product = "产品B" if product == "产品A" else "产品A"
        return {
            "status": "ok",
            "results": [
                {
                    "id": f"chunk-{product}",
                    "kb_id": kb_ids[0],
                    "file_id": f"file-{product}",
                    "content": f"{product} 的真实能力证据",
                    "metadata": {"source": f"{product}白皮书.docx"},
                },
                {
                    "id": f"chunk-{other_product}",
                    "kb_id": kb_ids[0],
                    "file_id": f"file-{other_product}",
                    "content": f"{other_product} 的其他能力证据",
                    "metadata": {"source": f"{other_product}白皮书.docx"},
                },
            ],
        }

    monkeypatch.setattr(industry_tools, "retrieve_kbs", fake_retrieve)
    result = await industry_tools.research_industry_products.coroutine(
        industry="智慧园区",
        requirement="统一管理终端",
        products=["产品A", "产品B"],
        kb_ids=["kb-1"],
        runtime=SimpleNamespace(context=SimpleNamespace()),
    )

    assert len(queries) == 2
    assert any("产品：产品A" in query for query in queries)
    assert any("产品：产品B" in query for query in queries)
    assert [item["product"] for item in result["products"]] == ["产品A", "产品B"]
    assert result["products"][0]["evidence"][0]["source"]["product"] == "产品A"
    assert result["products"][1]["evidence"][0]["source"]["product"] == "产品B"
    assert result["products"][0]["evidence"][0]["source"] == {
        "reference": 1,
        "product": "产品A",
        "title": "产品A白皮书.docx",
        "kb_id": "kb-1",
        "file_id": "file-产品A",
        "chunk_id": "chunk-产品A",
    }
    assert "url" not in result["products"][0]["evidence"][0]["source"]
    assert [len(item["evidence"]) for item in result["products"]] == [1, 1]
    assert [source["reference"] for source in result["sources"]] == [1, 2]


def test_industry_solution_request_accepts_empty_industry_and_products():
    # 前端「直接在输入框输入需求」模式：行业与产品可留空，仅需求必填
    request = IndustrySolutionRequest(industry="", requirement="统一管理终端", products=[])
    assert request.industry == ""
    assert request.products == []


def test_industry_solution_request_rejects_empty_requirement():
    with pytest.raises(ValidationError, match="需求不能为空"):
        IndustrySolutionRequest(industry="制造业", requirement="  ", products=[])


def test_industry_solution_request_rejects_too_many_products():
    with pytest.raises(ValidationError):
        IndustrySolutionRequest(
            industry="制造业",
            requirement="降本增效",
            products=[f"产品-{index}" for index in range(6)],
        )


def test_industry_solution_request_normalizes_products():
    request = IndustrySolutionRequest(
        industry=" 制造业 ", requirement=" 降本增效 ", products=[" 产品A ", "产品B", "产品b"]
    )
    assert request.industry == "制造业"
    assert request.requirement == "降本增效"
    assert request.products == ["产品A", "产品B"]


@pytest.mark.parametrize("products", [[], ["产品A"], ["产品A", " 产品A "]])
def test_product_research_input_rejects_less_than_two_products(products):
    # 工具执行层仍要求至少 2 个不同产品，避免检索无产品依据
    with pytest.raises(ValidationError):
        industry_tools.ProductResearchInput(
            industry="制造业",
            requirement="降本增效",
            products=products,
            kb_ids=["kb-1"],
        )


def test_structured_request_is_preserved_and_injected_into_model_context():
    request = {"industry": "智慧园区", "requirement": "统一管理", "products": ["产品A", "产品B"]}
    message = build_chat_input_message("生成行业方案", industry_solution=request)

    assert message.content == "生成行业方案"
    assert message.extra_metadata["industry_solution"] == request
    model_content = message.require_langchain_message().content
    assert "<industry_solution_request>" in model_content
    assert '"products": ["产品A", "产品B"]' in model_content


def test_agent_run_schema_validates_structured_products():
    payload = AgentRunCreate(
        query="生成行业方案",
        agent_slug="default",
        thread_id="thread-1",
        industry_solution={"industry": "能源", "requirement": "安全运营", "products": ["A", "B"]},
    )
    assert payload.industry_solution.products == ["A", "B"]


def test_render_solution_docx_contains_body_products_and_bilingual_references():
    document_bytes, chat_markdown = render_solution_docx(
        title="智慧园区行业解决方案",
        industry="智慧园区 / Smart Campus",
        products=["产品A", "Product B"],
        content="# 总体方案\n\n产品A 与 Product B 协同工作。[1][2]",
        sources=[
            {"reference": 1, "product": "产品A", "title": "产品A白皮书"},
            {"reference": 2, "product": "Product B", "title": "Product B Guide"},
        ],
    )
    document = Document(io.BytesIO(document_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert document_bytes.startswith(b"PK")
    assert "产品A 与 Product B 协同工作。[1][2]" in text
    assert "产品A, Product B" in text
    assert "[1] 产品A - 产品A白皮书" in text
    assert "[2] Product B - Product B Guide" in text
    assert chat_markdown.endswith("[2] Product B - `Product B Guide`")


def test_render_solution_docx_formats_inline_markdown_and_fonts_consistently():
    document_bytes, chat_markdown = render_solution_docx(
        title="智慧园区行业解决方案",
        industry="智慧园区",
        products=["Product A", "Product B", "Product C"],
        content=(
            "# 总体方案\n\n"
            "**安全管理**保障园区运行。[1]\n\n"
            "## 感知层\n\n"
            "- **运营管理**统一汇聚数据\n"
            "- **节能降耗**降低运营成本\n\n"
            "| 层级 | 能力 |\n| --- | --- |\n| 决策层 | **协同决策** |"
        ),
        sources=[
            {
                "reference": 1,
                "product": "Product A",
                "title": "product-a-smart-access.md",
                "kb_id": "kb-a",
                "file_id": "file-a",
                "chunk_id": "chunk-a",
            }
        ],
    )
    document = Document(io.BytesIO(document_bytes))
    paragraphs = list(document.paragraphs) + [
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    ]
    text = "\n".join(paragraph.text for paragraph in paragraphs)
    runs = [run for paragraph in paragraphs for run in paragraph.runs]

    assert "**" not in text
    assert "安全管理" in text
    assert "感知层" in text
    assert any(run.text == "安全管理" and run.bold for run in runs)
    assert any(run.text == "运营管理" and run.bold for run in runs)
    assert any(run.text == "协同决策" and run.bold for run in runs)
    assert "[1] Product A - product-a-smart-access.md" in text
    assert "`product-a-smart-access.md`" in chat_markdown
    assert "http://product-a-smart-access.md" not in chat_markdown

    # 字体对齐甲方（POCSTARS）样张：西文/数字 Arial，正文与标题中文宋体。
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = document.styles[style_name]
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        assert style.font.name == "Arial"
        assert fonts.get(qn("w:ascii")) == "Arial"
        assert fonts.get(qn("w:hAnsi")) == "Arial"
        assert fonts.get(qn("w:cs")) == "Arial"
        assert fonts.get(qn("w:eastAsia")) == "宋体"

    body_runs = [run for paragraph in document.paragraphs for run in paragraph.runs if run.text.strip()]
    table_runs = [
        run
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    ]
    assert body_runs
    assert table_runs

    def _assert_run_fonts(runs, cjk):
        for run in runs:
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            assert run.font.name == "Arial"
            assert fonts.get(qn("w:ascii")) == "Arial"
            assert fonts.get(qn("w:hAnsi")) == "Arial"
            assert fonts.get(qn("w:eastAsia")) == cjk

    _assert_run_fonts(body_runs, "宋体")
    _assert_run_fonts(table_runs, "黑体")


def test_render_solution_docx_rejects_missing_reference_and_omits_unused_sources():
    with pytest.raises(ValueError, match=r"不存在的来源.*\[3]"):
        render_solution_docx(
            title="方案",
            industry="制造业",
            products=["A", "B"],
            content="正文结论。[3]",
            sources=[{"reference": 1, "product": "A", "title": "A手册"}],
        )

    _, chat_markdown = render_solution_docx(
        title="方案",
        industry="制造业",
        products=["A", "B"],
        content="正文结论。[1]",
        sources=[
            {"reference": 1, "product": "A", "title": "A手册"},
            {"reference": 2, "product": "B", "title": "B手册"},
        ],
    )
    assert "[1] A - `A手册`" in chat_markdown
    assert "[2] B - B手册" not in chat_markdown


def test_render_solution_docx_replaces_model_generated_reference_list():
    _, chat_markdown = render_solution_docx(
        title="方案",
        industry="智慧园区",
        products=["A", "B"],
        content="正文结论。[1]\n\n## 来源 / References\n\n[1] 模型自行生成的来源",
        sources=[{"reference": 1, "product": "A", "title": "A手册"}],
    )

    assert chat_markdown.count("# 来源 / References") == 1
    assert "模型自行生成的来源" not in chat_markdown
    assert chat_markdown.endswith("[1] A - `A手册`")


def test_render_solution_docx_has_whitepaper_layout():
    document_bytes, _ = render_solution_docx(
        title="智慧园区行业解决方案",
        industry="智慧园区",
        products=["产品A", "产品B"],
        content=(
            "# 1. 总体方案\n\n产品协同能力说明。[1]\n\n"
            "## 1.1 系统架构\n\n| 层级 | 能力 |\n| --- | --- |\n| 决策层 | 统一决策 |"
        ),
        sources=[{"reference": 1, "product": "产品A", "title": "产品A手册"}],
    )
    document = Document(io.BytesIO(document_bytes))
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

    # 两节：封面 + 目录/正文
    assert len(document.sections) == 2
    # 纸型对齐甲方白皮书案例：A4 + 1.81/2.00cm 边距（新增节继承，检查首节）
    section = document.sections[0]
    assert section.page_width.twips == 11906  # 21.0cm
    assert section.page_height.twips == 16838  # 29.7cm
    assert section.left_margin.twips == 1026  # 1.81cm
    assert section.top_margin.twips == 1134  # 2.00cm
    # 封面含标题；正文以「方案概述」章节开头（标题不再单独出现在正文）
    document_text = "\n".join(paragraph_texts)
    assert "智慧园区行业解决方案" in document_text
    assert "方案概述" in document_text
    # POCSTARS 品牌：封面 logo + 公司名；版权信息/免责声明页；文末 AI 生成提示
    assert document.inline_shapes  # 封面 logo 图片
    assert "善理通益信息科技（深圳）有限公司" in document_text
    assert "版权信息" in document_text
    assert "免责声明" in document_text
    assert "请注意，本内容由AI生成。" in document_text
    # 目录域 + 打开时自动更新域
    assert any("TOC" in paragraph._p.xml for paragraph in document.paragraphs)
    assert "updateFields" in document.settings.element.xml
    # 页眉 = 方案标题，页脚含 PAGE 域
    header_text = document.sections[1].header.paragraphs[0].text
    assert header_text == "智慧园区行业解决方案"
    assert "PAGE" in document.sections[1].footer.paragraphs[0]._p.xml
    # 标题强调色
    heading1 = next(paragraph for paragraph in document.paragraphs if paragraph.style.name == "Heading 1")
    assert str(heading1.style.font.color.rgb) == "333399"
    # 表格表头 #333399 填充 + 白字加粗
    assert document.tables
    header_cell = document.tables[0].rows[0].cells[0]
    assert "333399" in header_cell._tc.xml
    assert header_cell.paragraphs[0].runs[0].bold
    assert str(header_cell.paragraphs[0].runs[0].font.color.rgb) == "FFFFFF"


def test_docx_filename_is_sanitized():
    filename = sanitize_docx_filename("../../智慧园区:方案\\报告")
    assert filename == "智慧园区-方案-报告.docx"
    assert ".." not in filename
    assert "/" not in filename
    assert "\\" not in filename


def test_export_source_rejects_internal_url_and_strips_path():
    source = industry_tools.SolutionSource(
        reference=1,
        product="A",
        title=r"C:\internal\A手册.docx",
        url="https://example.test/a",
    )
    assert source.title == "A手册.docx"

    with pytest.raises(ValidationError, match="HTTP/HTTPS"):
        industry_tools.SolutionSource(
            reference=1,
            product="A",
            title="A手册.docx",
            url="file:///internal/A手册.docx",
        )


def test_export_tool_writes_docx_inside_thread_outputs(tmp_path: Path, monkeypatch):
    outputs = tmp_path / "outputs"

    monkeypatch.setattr(industry_tools, "ensure_thread_dirs", lambda thread_id, uid: outputs.mkdir())
    monkeypatch.setattr(industry_tools, "sandbox_outputs_dir", lambda thread_id: outputs)
    monkeypatch.setattr(
        industry_tools,
        "virtual_path_for_thread_file",
        lambda thread_id, path, uid: f"/home/gem/user-data/outputs/{Path(path).name}",
    )

    result = industry_tools.export_industry_solution_docx.func(
        title="../智慧园区方案",
        industry="智慧园区",
        products=["产品A", "Product B"],
        content="# 总体方案\n\n协同能力。[1]",
        sources=[
            {
                "reference": 1,
                "product": "产品A",
                "title": "产品A手册",
                "kb_id": "kb-1",
                "file_id": "file-1",
                "chunk_id": "chunk-1",
            }
        ],
        runtime=SimpleNamespace(context=SimpleNamespace(thread_id="thread-1", uid="user-1")),
    )

    generated = outputs / Path(result["file_path"]).name
    assert generated.exists()
    assert generated.read_bytes().startswith(b"PK")
    assert ".." not in generated.name
    assert result["chat_markdown"].endswith("[1] 产品A - `产品A手册`")
