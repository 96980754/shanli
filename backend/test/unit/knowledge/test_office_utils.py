from __future__ import annotations

import io
import os
import tempfile

import pytest

from yuxi.knowledge.utils.office_writer import (
    markdown_to_blocks,
    markdown_to_sheets,
    serialize_edited_content,
    write_docx,
    write_xlsx,
)


def test_write_docx_roundtrip():
    blocks = [
        {"kind": "heading", "text": "产品介绍"},
        {"kind": "para", "text": "这是正文内容。"},
        {"kind": "table", "rows": [["参数", "值"], ["电池", "5000mAh"]]},
    ]
    data = write_docx(blocks)
    assert len(data) > 0

    # 写回临时文件再解析，验证内容
    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(data)
        tmp = f.name
    try:
        doc = Document(tmp)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "产品介绍" in texts
        assert "这是正文内容。" in texts
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "参数"
    finally:
        os.unlink(tmp)


def test_write_xlsx_roundtrip():
    sheets = [{"name": "Sheet1", "rows": [["参数", "值"], ["电池", "5000mAh"]]}]
    data = write_xlsx(sheets)
    assert len(data) > 0

    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True)
    ws = wb.active
    assert ws.title == "Sheet1"
    assert ws.cell(1, 1).value == "参数"
    assert ws.cell(2, 2).value == "5000mAh"


def test_markdown_to_blocks_parses_headings_paras_and_tables():
    md = "# 产品介绍\n\n一段正文。\n\n| 参数 | 值 |\n| --- | --- |\n| 电池 | 5000mAh |\n"
    blocks = markdown_to_blocks(md)

    kinds = [b["kind"] for b in blocks]
    assert kinds == ["heading", "para", "table"]
    assert blocks[0]["text"] == "产品介绍"
    assert blocks[0]["level"] == 1
    assert blocks[1]["text"] == "一段正文。"
    assert blocks[2]["rows"] == [["参数", "值"], ["电池", "5000mAh"]]


def test_markdown_to_blocks_handles_multiple_headings_levels():
    md = "# 一级\n\n## 二级\n\n正文\n"
    blocks = markdown_to_blocks(md)
    assert [(b["kind"], b.get("text"), b.get("level")) for b in blocks] == [
        ("heading", "一级", 1),
        ("heading", "二级", 2),
        ("para", "正文", None),
    ]


def test_markdown_to_blocks_empty_returns_empty():
    assert markdown_to_blocks("") == []
    assert markdown_to_blocks("   \n  ") == []


def test_markdown_to_sheets_puts_text_and_tables_into_rows():
    md = "# 标题\n\n| 参数 | 值 |\n| --- | --- |\n| 电池 | 5000mAh |\n"
    sheets = markdown_to_sheets(md)

    assert len(sheets) == 1
    assert sheets[0]["name"] == "Sheet1"
    assert sheets[0]["rows"] == [
        ["标题"],
        ["参数", "值"],
        ["电池", "5000mAh"],
    ]


def test_markdown_to_sheets_empty_returns_empty():
    assert markdown_to_sheets("") == []
    assert markdown_to_sheets("   \n  ") == []


def test_markdown_to_blocks_write_docx_roundtrip():
    md = "# 产品介绍\n\n正文内容\n\n| 参数 | 值 |\n| --- | --- |\n| 电池 | 5000mAh |\n"
    data = write_docx(markdown_to_blocks(md))
    assert len(data) > 0

    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(data)
        tmp = f.name
    try:
        doc = Document(tmp)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "产品介绍" in texts
        assert "正文内容" in texts
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(1, 1).text == "5000mAh"
    finally:
        os.unlink(tmp)


def test_serialize_edited_content_by_type():
    docx_data = serialize_edited_content("docx", {"blocks": [{"kind": "para", "text": "x"}]})
    assert len(docx_data) > 0
    xlsx_data = serialize_edited_content("xlsx", {"sheets": [{"name": "S", "rows": [["a"]]}]})
    assert len(xlsx_data) > 0


@pytest.mark.asyncio
async def test_extract_office_content_docx(monkeypatch):
    from yuxi.knowledge import runtime as kb_runtime

    blocks = [
        {"kind": "heading", "text": "标题"},
        {"kind": "para", "text": "段落"},
    ]
    docx_bytes = write_docx(blocks)

    class _Kb:
        async def _read_minio_bytes(self, file_path):
            return docx_bytes

    class _Manager:
        async def _get_kb_for_database(self, kb_id):
            return _Kb()

    monkeypatch.setattr(kb_runtime, "knowledge_base", _Manager())

    from yuxi.knowledge.utils.office_content import extract_office_content

    result = await extract_office_content("kb-1", "minio://kb/x.docx", "x.docx")
    assert result["type"] == "docx"
    assert any(b.get("text") == "标题" for b in result["blocks"])
    assert any(b.get("text") == "段落" for b in result["blocks"])
